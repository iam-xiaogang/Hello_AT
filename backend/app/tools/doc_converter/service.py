"""Document conversion service (PDF ↔ Word / text).

Supported targets:
- ``pdf-to-word``  — layout-preserving PDF → DOCX via pdf2docx
- ``pdf-to-text``  — plain-text extraction from PDF via PyMuPDF
- ``word-to-text`` — plain-text extraction from DOCX via python-docx
"""
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException, status
from pymupdf import Document as PdfDocument

MAX_UPLOAD_BYTES = 50 * 1024 * 1024

TARGETS = {"pdf-to-word", "pdf-to-text", "word-to-text"}
SOURCE_EXTENSIONS = {"pdf-to-word": "pdf", "pdf-to-text": "pdf", "word-to-text": "docx"}

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MIME = "text/plain; charset=utf-8"

OUTPUT_EXTENSION = {"pdf-to-word": "docx", "pdf-to-text": "txt", "word-to-text": "txt"}


def _stem(filename: str | None) -> str:
    return (filename or "document").rsplit(".", 1)[0]


def _convert_pdf_to_word(raw: bytes, filename: str | None) -> tuple[bytes, str, str]:
    from pdf2docx import Converter  # imported lazily so other targets need no extra deps at call time

    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = Path(tmp) / "input.pdf"
        docx_path = Path(tmp) / "output.docx"
        pdf_path.write_bytes(raw)
        try:
            converter = Converter(str(pdf_path))
            try:
                converter.convert(str(docx_path), start=0, end=None)
            finally:
                converter.close()
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="无法转换该 PDF（可能已加密、损坏或包含扫描页）。",
            ) from exc
        content = docx_path.read_bytes()
    return content, DOCX_MIME, f"{_stem(filename)}.docx"


def _convert_pdf_to_text(raw: bytes, filename: str | None) -> tuple[bytes, str, str]:
    try:
        document = PdfDocument(stream=raw, filetype="pdf")
    except Exception as exc:
        raise HTTPException(status_code=400, detail="无法读取该 PDF 文件。") from exc
    if document.needs_pass:
        document.close()
        raise HTTPException(status_code=400, detail="该 PDF 已加密，无法转换。")
    try:
        pages = [page.get_text() for page in document]
    finally:
        document.close()
    text = "\n\n".join(page.strip() for page in pages if page.strip())
    return text.encode("utf-8"), TEXT_MIME, f"{_stem(filename)}.txt"


def _convert_word_to_text(raw: bytes, filename: str | None) -> tuple[bytes, str, str]:
    try:
        document = Document(BytesIO(raw))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="无法读取该 Word 文档。") from exc
    parts: list[str] = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return text.encode("utf-8"), TEXT_MIME, f"{_stem(filename)}.txt"


def convert_document(raw: bytes, filename: str | None, target: str) -> tuple[bytes, str, str]:
    """Validate inputs and run the requested conversion (call in a worker thread)."""
    if target not in TARGETS:
        raise HTTPException(status_code=400, detail="不支持的目标格式。")
    if not raw:
        raise HTTPException(status_code=400, detail="上传的文件为空。")

    source = SOURCE_EXTENSIONS[target]
    if filename and not filename.lower().endswith(f".{source}"):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"目标格式「{target}」需要上传 {source.upper()} 文件。",
        )

    if target == "pdf-to-word":
        return _convert_pdf_to_word(raw, filename)
    if target == "pdf-to-text":
        return _convert_pdf_to_text(raw, filename)
    return _convert_word_to_text(raw, filename)
